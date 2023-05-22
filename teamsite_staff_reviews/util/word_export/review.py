import logging

from django.db.models import Max
from django.utils.timezone import now
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from reviews.models import Nomination, ReviewForm, ReviewFormResponse
from reviews.util.word_export.util import add_markdown, add_page_number, logo

logger = logging.getLogger(__name__)


class ReviewDocument:
    def __init__(self, *nominations):
        self.__document = Document()
        self.__nominations = 0
        self.__init_doc()
        for nom in nominations:
            try:
                self.add_nomination_to_doc(nom)
            except:
                logger.exception("Failed to add %s to word export")

    def __init_doc(self):
        document = self.__document
        header = document.sections[0].header.paragraphs[0]
        logo_run = header.add_run()
        logo_run.add_picture(str(logo.absolute()), width=Inches(1))
        header.add_run().text = "\t\tConfidential"

        footer = document.sections[0].footer.paragraphs[0]
        footer.add_run().text = f"Created {now():%d %b %Y %H:%M}\t\tPage "

        add_page_number(footer.add_run(), "PAGE")
        footer.add_run(" of ")
        add_page_number(footer.add_run(), "NUMPAGES")

        style = document.styles.add_style("Normal Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.size = Pt(8)
        style.font.italic = True

        style = document.styles.add_style("List Bullet Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["List Bullet"]
        style.font.size = Pt(8)
        style.font.italic = True

    def add_nomination_to_doc(self, nomination: Nomination):
        document = self.__document
        self.__nominations += 1
        if self.__nominations > 1:
            document.add_page_break()

        form = ReviewForm.objects.filter(
            period=nomination.period, role=nomination.role
        ).first()
        questions = form.questions
        reviewee = nomination.reviewee

        try:
            invitation = nomination.invitation
        except:
            invitation = None

        if invitation and nomination.invitation.user:
            reviewer = (
                f"{nomination.external_name} ({nomination.invitation.user.email})"
            )
        elif nomination.reviewer is None:
            reviewer = f"{nomination.external_name}"
        else:
            reviewer = (
                f"{nomination.reviewer.first_name} {nomination.reviewer.last_name}"
            )

        document.add_heading(
            f"{reviewee.first_name} {reviewee.last_name}\n"
            f"{form.period.year} {form.period.round_label} {form.title}",
            0,
        )
        p = document.add_paragraph()
        p.add_run(f"Reviewer: {reviewer}", "Book Title")

        last_mod = (
            ReviewFormResponse.objects.filter(nomination=nomination)
            .values("last_modified")
            .aggregate(last_mod=Max("last_modified"))
            .get("last_mod")
        )

        p = document.add_paragraph()
        last_mod = f"{last_mod:%d %b %Y %H:%M}" if last_mod else "Never"
        p.add_run(f"Last Modified: {last_mod}", "Book Title")

        add_markdown(document, form.description)

        for q in questions.all():
            document.add_heading(q.title, level=1)
            add_markdown(document, q.description, style_name="Small")

            response = q.responses.filter(nomination=nomination).first()
            if response is not None:
                add_markdown(document, response.value)

    def save(self, filename):
        self.__document.save(filename)
