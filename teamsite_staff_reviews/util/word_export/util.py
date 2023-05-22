from pathlib import Path

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx
from marko import Markdown

logo = Path(__file__).parent / "../../static/SF-logo-500x104.png"


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(run, code="PAGE"):
    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "begin")

    instrText = create_element("w:instrText")
    create_attribute(instrText, "xml:space", "preserve")
    instrText.text = code

    fldChar2 = create_element("w:fldChar")
    create_attribute(fldChar2, "w:fldCharType", "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_markdown(document, text, style_name=None):
    if text is None:
        document.add_paragraph("")
        return

    doc_length = len(document.paragraphs)

    markdown = Markdown(extensions=[QuoteRenderer])
    html = markdown.convert(text)
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(html, document)

    if style_name:
        for p in document.paragraphs[doc_length:]:
            p.style = f"{p.style.name} {style_name}"


class QuoteRendererMixin:
    def render_quote(self, element):
        rendered = self.render_children(element)
        rendered = rendered.replace(
            "<p>", '<p style="margin-left: 15px;"><span style="color: #666666">'
        )
        rendered = rendered.replace("</p>", "</span></p>")
        return rendered


class QuoteRenderer:
    renderer_mixins = [QuoteRendererMixin]
