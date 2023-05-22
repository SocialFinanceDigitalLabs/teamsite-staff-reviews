from datetime import date, timedelta

from django.db.models import Max, Q, Sum
from django.utils.timezone import now
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from humanize import naturaldelta
from resourcing.models import EmploymentEvent, EmploymentEventType, TrackedTime
from reviews.models import Nomination, ReviewForm, ReviewFormResponse, ReviewRound
from reviews.util.word_export.util import add_markdown, add_page_number, logo


class AssessmentDocument:
    def __init__(self, nomination, exclude_questions=None):
        self.document = Document()
        self.__init_doc()
        self.add_nomination_to_doc(nomination, exclude_questions=exclude_questions)

    def __init_doc(self):
        document = self.document
        header = document.sections[0].header.paragraphs[0]
        logo_run = header.add_run()
        logo_run.add_picture(str(logo.absolute()), width=Inches(1))
        header.add_run().text = "\t\tConfidential"

        footer = document.sections[0].footer.paragraphs[0]
        footer.add_run().text = f"Created {now():%d %b %Y %H:%M}\t\tPage "

        add_page_number(footer.add_run(), "PAGE")
        footer.add_run(" of ")
        add_page_number(footer.add_run(), "NUMPAGES")

        style = document.styles.add_style("Normal Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.size = Pt(8)
        style.font.italic = True

        style = document.styles.add_style("List Bullet Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["List Bullet"]
        style.font.size = Pt(8)
        style.font.italic = True

    def add_nomination_to_doc(self, nomination: Nomination, exclude_questions=None):
        document = self.document
        form = ReviewForm.objects.filter(
            period=nomination.period, role=nomination.role
        ).first()
        questions = form.questions
        if exclude_questions:
            questions = questions.filter(~Q(pk__in=exclude_questions))
        reviewee = nomination.reviewee

        try:
            invitation = nomination.invitation
        except:
            invitation = None

        if invitation and nomination.invitation.user:
            reviewer = (
                f"{nomination.external_name} ({nomination.invitation.user.email})"
            )
        elif nomination.reviewer is None:
            reviewer = f"{nomination.external_name}"
        else:
            reviewer = (
                f"{nomination.reviewer.first_name} {nomination.reviewer.last_name}"
            )

        document.add_heading(
            f"Appraisal Form - {form.period.year} {form.period.round_label}", 0
        )

        title_style = "Book Title"

        table = document.add_table(rows=2, cols=4)
        table.rows[0].cells[0].paragraphs[0].add_run("Line report", title_style)
        table.rows[0].cells[1].text = f"{reviewee.first_name} {reviewee.last_name}"
        table.rows[0].cells[2].paragraphs[0].add_run("Line manager", title_style)
        table.rows[0].cells[3].text = reviewer
        table.rows[1].cells[0].paragraphs[0].add_run("Current team", title_style)

        team = reviewee.teams.first()
        if team:
            table.rows[1].cells[1].text = f"{team.team.name}"

        last_mod = (
            ReviewFormResponse.objects.filter(nomination=nomination)
            .values("last_modified")
            .aggregate(last_mod=Max("last_modified"))
            .get("last_mod")
        )
        last_mod = f"{last_mod:%d %b %Y %H:%M}" if last_mod else ""

        table.rows[1].cells[2].paragraphs[0].add_run("Completion date", title_style)
        table.rows[1].cells[3].text = last_mod

        if nomination.period.round == ReviewRound.MID_YEAR:
            period_start = date(nomination.period.year, 1, 1)
            period_end = date(nomination.period.year, 6, 30)
        else:
            period_start = date(nomination.period.year, 7, 1)
            period_end = date(nomination.period.year, 12, 31)

        projects = (
            TrackedTime.objects.filter(
                ~Q(project__team__name="Internal"),
                user=reviewee,
                week__gte=period_start,
                week__lte=period_end,
                time__gt=0,
            )
            .values(
                "project__survey_form",
            )
            .annotate(time_sum=Sum("time"))
            .filter(time_sum__gte=3)
            .order_by("-time_sum")
        )

        projects = ", ".join(
            [f"{p['project__survey_form']} ({p['time_sum']})" for p in projects]
        )

        para = document.add_paragraph()
        para.add_run("Projects in period", title_style)
        para.add_run("\n")
        para.add_run(projects)

        document.add_heading("Career History at Social Finance", 1)

        joined = (
            EmploymentEvent.objects.filter(
                user=reviewee, event_type=EmploymentEventType.JOINED
            )
            .order_by("date")
            .last()
        )

        promoted = (
            EmploymentEvent.objects.filter(
                user=reviewee, event_type=EmploymentEventType.PROMOTED
            )
            .order_by("date")
            .last()
        )

        table = document.add_table(rows=2, cols=4)
        table.rows[0].cells[0].paragraphs[0].add_run("Start date", title_style)
        table.rows[0].cells[1].text = (
            f"{joined.date:%d %b %Y}\n({naturaldelta(now().date() - joined.date)})"
            if joined
            else ""
        )
        table.rows[1].cells[0].paragraphs[0].add_run(
            "Date of last promotion", title_style
        )
        table.rows[1].cells[1].text = (
            f"{promoted.date:%d %b %Y}\n({naturaldelta(now().date() - promoted.date)})"
            if promoted
            else ""
        )

        table.rows[0].cells[2].paragraphs[0].add_run("Joining cohort", title_style)
        table.rows[0].cells[3].text = (
            joined.cohort.name if joined and joined.cohort else ""
        )
        table.rows[1].cells[2].paragraphs[0].add_run("Current cohort", title_style)
        table.rows[1].cells[3].text = f"{reviewee.profile.cohort.name}"

        add_markdown(document, form.description)

        for q in questions.all():
            document.add_heading(q.title, level=1)
            # add_markdown(document, q.description, style_name='Small')
            response = q.responses.filter(nomination=nomination).first()
            if response is not None:
                add_markdown(document, response.value.strip())

    def save(self, filename):
        self.document.save(filename)
